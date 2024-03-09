import streamlit as st
import json
import random
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from datetime import datetime
import base64
import io
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
st.set_page_config(
    page_title="Islamic  Style Prompts Generator",
    page_icon='⚡️',
    layout="wide"
)
import style_functions as sf
sf.hide_all()
#sf.button_color()
#sf.footer_style_custom_position()
sf.custom_style()
def load_categories_from_json(file_path):
    with open(file_path, 'r') as file:
        categories = json.load(file)
    return categories

# Function to generate random choice from a category
def generate_random_item(category):
    if category in categories:
        items = categories[category]["random"]
        return random.choice(items)
    else:
        return None

#documnet
def document_sd(prompts):

    doc = Document()
    # Set margins for the document (in inches)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Title
    title = doc.add_heading("Prompts", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Header
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Header - Your Company Name"

    # Footer
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Footer - Copyright © Your Company"

    # Adding numbered prompts manually
    for i, prompt in enumerate(prompts, start=1):
        para = doc.add_paragraph()
        #para.add_run(f"{i}. ").bold = True
        #para.add_run(f"{i}. ").bold = True
        para.add_run(prompt)

    # Save the document
    #doc.save("prompts.docx")
    stamp_time = datetime.now().strftime("%Y%m%d%H%M%S")
    docx_file = f"{stamp_time}_SD_Prompts.docx"
    docx_content = io.BytesIO()
    doc.save(docx_content)
    docx_content.seek(0)
        
    st.download_button('Download file', docx_content, key='download_button2', file_name=f"{docx_file}", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.write(f"Prompts saved to prompts.docx: {len(prompts)}")

# Streamlit UI
st.title("Choose Items")

json_file_path = "islamic_style.json"
categories = load_categories_from_json(json_file_path)

category1 = st.selectbox("Select Type of Photography", categories["Type of photography"]["random"])
category2 = st.selectbox("Select Islamic 3D Frames", categories["Islamic 3D Frames"]["random"])
category3 = st.selectbox("Select Lights", categories["Lights"]["random"])
category4 = st.selectbox("Select Camera Type", categories["Camera Type"]["random"])
category5 = st.selectbox("Select Camera Angles", categories["Camera Angles"]["random"])
category6 = st.selectbox("Select Environments", categories["Environments"]["random"])
category7 = st.selectbox("Select Colors", categories["Colors"]["random"])
category8 = st.selectbox("Select Background", categories["Background"]["random"])
category9 = st.selectbox("Select Actions", categories["Actions"]["random"])
category10 = st.selectbox("Select Frame Shapes", categories["Frame Shapes"]["random"])
category11 = st.selectbox("Select Materials", categories["Materials"]["random"])
category12 = st.selectbox("Select Subjects", categories["Subjects"]["random"])



choice_type = st.radio("Select choice type", ("Specific", "Random"))

num_prompts = st.number_input("Enter the number of prompts to generate", min_value=1, step=1, value=1)

prompts = []

if category1 and category2 and category3:
    for i in range(num_prompts):
        if choice_type == "Specific":
            prompts.append(f"{i+1}. Prompt: {category1}, {category2}, {category3}, {category4}, {category5}, {category6}, {category7}, {category8}, {category9}, {category10}, {category11}, {category12}")
        elif choice_type == "Random":
            
            random_category1 = generate_random_item("Type of photography")
            random_category2 = generate_random_item("Islamic 3D Frames")
            random_category3 = generate_random_item("Lights")
            random_category4 = generate_random_item("Camera Type")
            random_category5 = generate_random_item("Camera Angles")
            random_category6 = generate_random_item("Environments")
            random_category7 = generate_random_item("Colors")
            random_category8 = generate_random_item("Background")
            random_category9 = generate_random_item("Actions")
            random_category10 = generate_random_item("Frame Shapes")
            random_category11 = generate_random_item("Materials")
            random_category12 = generate_random_item("Subjects")
            

            
                
            prompts.append(f" a large frame attached to the wall of a {random_category1}, {random_category2}, {random_category3}, {random_category4}, {random_category5}, {random_category6}, {random_category7}, {random_category8}, {random_category9}, {random_category10}, {random_category11}, {random_category12}" + " jumping out of an ornate gilded picture frame and into the physical world running outside of the frame towards the museum, he patrons of the museum are fleeing,  real painting, natural lighting, manipulation")
            #prompts.append(f": {random_category1}, {random_category2}, {random_category3}, {random_category4}, {random_category5}, {random_category6}, {random_category7}, {random_category8}, {random_category9}, {random_category10}, {random_category11}, {random_category12}, {random_category13}, {random_category14}, {random_category15}, {random_category16}, {random_category17}, {random_category18}, {random_category19}, {random_category20}, {random_category21}, {random_category22}, {random_category23}, {random_category24}, {random_category25}, {random_category26}, {random_category27}, {random_category28}, {random_category29}, {random_category30}")

document_sd(prompts)
# Save prompts to Word document
